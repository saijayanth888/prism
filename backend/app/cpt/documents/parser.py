"""Document parsers — extract raw text from multiple file formats.

Supported formats (optional deps degrade gracefully):
  .docx → python-docx   | .xlsx → openpyxl
  .pdf  → pypdf          | .yaml/.json → stdlib
  .md   → stdlib         | .csv → stdlib
  .txt  → stdlib
"""
from __future__ import annotations

import csv
import io
import json
import logging
import re
from abc import ABC, abstractmethod
from pathlib import Path

log = logging.getLogger(__name__)


class ParsedDocument:
    def __init__(self, text: str, metadata: dict | None = None):
        self.text = text
        self.metadata: dict = metadata or {}
        self.page_count: int = 1
        self.word_count: int = len(text.split())


class BaseParser(ABC):
    @abstractmethod
    def can_parse(self, mime_type: str, extension: str) -> bool: ...

    @abstractmethod
    def parse(self, content: bytes, filename: str) -> ParsedDocument: ...


class TextParser(BaseParser):
    def can_parse(self, mime_type: str, extension: str) -> bool:
        return extension in {".txt", ".log", ".env"} or mime_type.startswith("text/plain")

    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        text = content.decode("utf-8", errors="replace")
        return ParsedDocument(text, {"format": "text"})


class MarkdownParser(BaseParser):
    def can_parse(self, mime_type: str, extension: str) -> bool:
        return extension in {".md", ".mdx", ".markdown"}

    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        text = content.decode("utf-8", errors="replace")
        # Strip markdown syntax for cleaner extraction
        text_clean = re.sub(r"```[\s\S]*?```", " [CODE_BLOCK] ", text)
        text_clean = re.sub(r"`[^`]+`", " [INLINE_CODE] ", text_clean)
        text_clean = re.sub(r"#{1,6}\s+", "", text_clean)
        text_clean = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text_clean)
        return ParsedDocument(text_clean, {"format": "markdown", "original": text})


class JsonParser(BaseParser):
    def can_parse(self, mime_type: str, extension: str) -> bool:
        return extension in {".json", ".jsonl"} or "json" in mime_type

    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        try:
            data = json.loads(content.decode("utf-8", errors="replace"))
            text = json.dumps(data, indent=2)
            return ParsedDocument(text, {"format": "json"})
        except Exception:
            return ParsedDocument(content.decode("utf-8", errors="replace"), {"format": "json", "parse_error": True})


class YamlParser(BaseParser):
    def can_parse(self, mime_type: str, extension: str) -> bool:
        return extension in {".yaml", ".yml"}

    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        try:
            import yaml  # type: ignore
            data = yaml.safe_load(content.decode("utf-8", errors="replace"))
            text = yaml.dump(data, default_flow_style=False) if data else ""
            return ParsedDocument(text, {"format": "yaml"})
        except ImportError:
            return ParsedDocument(content.decode("utf-8", errors="replace"), {"format": "yaml"})
        except Exception as e:
            return ParsedDocument(content.decode("utf-8", errors="replace"), {"format": "yaml", "parse_error": str(e)})


class CsvParser(BaseParser):
    def can_parse(self, mime_type: str, extension: str) -> bool:
        return extension == ".csv" or "csv" in mime_type

    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        text_io = io.StringIO(content.decode("utf-8", errors="replace"))
        reader = csv.reader(text_io)
        rows = list(reader)
        if not rows:
            return ParsedDocument("", {"format": "csv", "rows": 0})
        headers = rows[0]
        lines = [", ".join(headers)]
        for row in rows[1:]:
            lines.append(", ".join(f"{h}={v}" for h, v in zip(headers, row) if v))
        return ParsedDocument("\n".join(lines), {"format": "csv", "rows": len(rows), "columns": len(headers)})


class WordParser(BaseParser):
    def can_parse(self, mime_type: str, extension: str) -> bool:
        return extension in {".docx", ".doc"} or "wordprocessingml" in mime_type

    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        try:
            import docx  # type: ignore
            doc = docx.Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Extract tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    table_texts.append(" | ".join(cell.text for cell in row.cells if cell.text.strip()))
            all_text = "\n".join(paragraphs + table_texts)
            return ParsedDocument(all_text, {"format": "docx", "paragraph_count": len(paragraphs)})
        except ImportError:
            log.warning("python-docx not installed, returning raw bytes as text")
            return ParsedDocument(content.decode("utf-8", errors="replace"), {"format": "docx", "parse_error": "missing python-docx"})
        except Exception as e:
            return ParsedDocument("", {"format": "docx", "parse_error": str(e)})


class ExcelParser(BaseParser):
    def can_parse(self, mime_type: str, extension: str) -> bool:
        return extension in {".xlsx", ".xls"} or "spreadsheetml" in mime_type

    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        try:
            import openpyxl  # type: ignore
            wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
            lines = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                lines.append(f"Sheet: {sheet_name}")
                for row in ws.iter_rows(values_only=True):
                    row_text = " | ".join(str(v) for v in row if v is not None)
                    if row_text:
                        lines.append(row_text)
            return ParsedDocument("\n".join(lines), {"format": "xlsx", "sheets": wb.sheetnames})
        except ImportError:
            return ParsedDocument("", {"format": "xlsx", "parse_error": "missing openpyxl"})
        except Exception as e:
            return ParsedDocument("", {"format": "xlsx", "parse_error": str(e)})


class PdfParser(BaseParser):
    def can_parse(self, mime_type: str, extension: str) -> bool:
        return extension == ".pdf" or "pdf" in mime_type

    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        # Try pypdf first, then pdfplumber, then fallback
        try:
            from pypdf import PdfReader  # type: ignore
            reader = PdfReader(io.BytesIO(content))
            pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n\n".join(pages)
            return ParsedDocument(text, {"format": "pdf", "page_count": len(pages)})
        except ImportError:
            pass
        try:
            import pdfplumber  # type: ignore
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages = [p.extract_text() or "" for p in pdf.pages]
            text = "\n\n".join(pages)
            return ParsedDocument(text, {"format": "pdf", "page_count": len(pages)})
        except ImportError:
            pass
        log.warning("No PDF parser installed (pypdf or pdfplumber); skipping")
        return ParsedDocument("", {"format": "pdf", "parse_error": "no_pdf_parser"})


class ParserFactory:
    """Select and run the right parser for a given file."""

    _parsers: list[BaseParser] = [
        WordParser(),
        ExcelParser(),
        PdfParser(),
        YamlParser(),
        JsonParser(),
        MarkdownParser(),
        CsvParser(),
        TextParser(),
    ]

    @classmethod
    def parse(cls, content: bytes, filename: str, mime_type: str = "") -> ParsedDocument:
        ext = Path(filename).suffix.lower()
        for parser in cls._parsers:
            if parser.can_parse(mime_type, ext):
                try:
                    return parser.parse(content, filename)
                except Exception as e:
                    log.error("parser_error", parser=type(parser).__name__, filename=filename, error=str(e))
        # Ultimate fallback — decode as text
        return ParsedDocument(content.decode("utf-8", errors="replace"), {"format": "unknown"})

    @classmethod
    def supported_extensions(cls) -> list[str]:
        return [".docx", ".doc", ".xlsx", ".xls", ".pdf", ".yaml", ".yml",
                ".json", ".jsonl", ".md", ".mdx", ".csv", ".txt", ".log"]
